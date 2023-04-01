#!/usr/bin/env python
# coding: utf-8

# In[ ]:


import os
import re
import docx
import PyPDF2
import io 
import nltk
import altair as alt
import numpy as np
nltk.download('stopwords')
nltk.download('punkt')
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize 
import pandas as pd
import streamlit as st
import numpy as np
# loading the trained model
#pickle_in = open('classifier.pkl', 'rb') 
#classifier = pickle.load(pickle_in)



#st.sidebar.subheader('File_Description')

html_temp3 = """
    <div style="background-color:black;padding:7px">
    <h2 style="color:white;text-align:center;"> RESUME SCREENING</h2>
    </div>
    """
st.markdown(html_temp3,unsafe_allow_html=True)
st.subheader(' ')

# Define key terms dictionary for fixing Role Applied for 
terms = {'WorkDay ERP':['workday', 'workday consultant', 'workday hcm', 'eib', 'picof', 
                        'workday studio','nnbound/outbound integrations'],
         'Peoplesoft':['peoplesoft', 'pia','ccb','birt','peci','ccw','pum','people tools',
                        'peoplesoft implementation','peoplesoft components',
                        'peoplesoft dba','peoplesoft admin','peoplesoft admin/dba','peopleSoft fscm', 
                        'peopletoolsupgrade','peopletools upgrade','process scheduler servers',
                        'peoplesoft hrms','peopleSoft consultant','peopledoft cloud',
                        'PeopleSoft migrations','eoplesoft Testing Framework','pure internet architecture'],             
         'Database Developer':['sql','sql server', 'ms sql server','msbi', 'sql developer', 'ssis','ssrs',
                        'ssms','t-sql','tsql','Razorsql', 'razor sql','triggers','powerbi','power bi',
                        'oracle sql', 'pl/sql', 'pl\sql','oracle', 'oracle 11g','oledb','cte','ddl',
                        'dml','etl','mariadb','maria db'],
         'Java Developer':['reactjs', 'react js', 'react js developer', 'html', 
                        'css3','xml','javascript','html5','boostrap','jquery', 'redux','php', 'node js',
                        'nodejs','apache','netbeans','nestjs','nest js','react developer','react hooks',
                        'jenkins']}

# List of all key terms to indicate skillset. Include all the key words in the list 
allTerms = ['workday', 'hcm', 'eib', 'picof','workday hcm',
                        'workday studio','nnbound/outbound integrations',
                        'peoplesoft', 'pia','ccb','birt','peci','ccw','pum','people tools',
                        'peoplesoft implementation','peoplesoft components',
                        'peoplesoft dba','peoplesoft admin','peoplesoft admin/dba','peopleSoft fscm', 
                        'peopletoolsupgrade','peopletools upgrade','process scheduler servers',
                        'peoplesoft hrms','peopleSoft consultant','peopledoft cloud',
                        'PeopleSoft migrations','eoplesoft Testing Framework','pure internet architecture',
                        'sql','sql server', 'ms sql server','msbi', 'sql developer', 'ssis','ssrs',
                        'ssms','t-sql','tsql','Razorsql', 'razor sql','triggers','powerbi','power bi',
                        'oracle sql', 'pl/sql', 'pl\sql','oracle', 'oracle 11g','oledb','cte','ddl',
                        'dml','etl','mariadb','maria db','reactjs', 'react js', 'react js developer', 'html', 
                        'css3','xml','javascript','html5','boostrap','jquery', 'redux','php', 'node js',
                        'nodejs','apache','netbeans','nestjs','nest js','react developer','react hooks',
                        'jenkins']

# Function to extract text from resume
def getText(filename):
      
    # Create empty string 
    fullText = ''
    if filename.endswith('.docx'):
        doc = docx.Document(filename)
        
        for para in doc.paragraphs:
            fullText = fullText + para.text
            
           
    elif filename.endswith('.pdf'):  
        with open(filename, "rb") as pdf_file:
            pdoc = PyPDF2.PdfFileReader(filename)
            number_of_pages = pdoc.getNumPages()
            page = pdoc.pages[0]
            page_content = page.extractText()
             
        for paragraph in page_content:
            fullText =  fullText + paragraph
            
    else:
        import aspose.words as aw
        output = aw.Document()
        # Remove all content from the destination document before appending.
        output.remove_all_children()
        input = aw.Document(filename)
        # Append the source document to the end of the destination document.
        output.append_document(input, aw.ImportFormatMode.KEEP_SOURCE_FORMATTING)
        output.save("Output.docx");
        doc = docx.Document('Output.docx')
        
        for para in doc.paragraphs:
            fullText = fullText + para.text
        fullText = fullText[79:]
         
    return (fullText)

# Function to remove punctuation and tokenize the text
def tokenText(extText):
   
    # Remove punctuation marks
    punc = '''!()-[]{};:'"\,.<>/?@#$%^&*_~'''
    for ele in extText:
        if ele in punc:
            puncText = extText.replace(ele, "")
            
    # Tokenize the text and remove stop words
    stop_words = set(stopwords.words('english'))
    puncText.split()
    word_tokens = word_tokenize(puncText)
    TokenizedText = [w for w in word_tokens if not w.lower() in stop_words]
    TokenizedText = []
  
    for w in word_tokens:
        if w not in stop_words:
            TokenizedText.append(w)
    return(TokenizedText)

# Function to read the tokenized text and search for the key words to dermine the Role Applied for
def roleApplied (Text):
    
    # covert the text to lower case
    for i in range(len(Text)):
        Text[i] = Text[i].lower()
    
    # Obtain the scores for each area
    for area in terms.keys():
        if area == 'WorkDay ERP':
            for word in terms[area]:
                if word in Text:
                    role = area
                    return (role)
                
        elif area == 'Peoplesoft':
            for word in terms[area]:
                if word in Text:
                    role = area
                    return(role)   
                
        elif area == 'Database Developer':
            for word in terms[area]:
                if word in Text:
                    role =  area
                    return(role)
            
        elif area == 'Java Developer':
             for word in terms[area]:
                if word in Text:
                    role = area
                    return(role)
        else:
            role = "Fresher"
            return(role)
        
# Function to extract Name and contact details
def contactDetails(Text):
    name = ''  
    for i in range(0,3):
        name = " ".join([name, Text[i]])
    return(name)

# Function to extract experience details
def expDetails(Text):
    global sent
   
    Text = Text.split()
   
    for i in range(len(Text)-2):
        Text[i].lower()
        
        if Text[i] ==  'years':
            sent =  Text[i-2] + ' ' + Text[i-1] +' ' + Text[i] +' '+ Text[i+1] +' ' + Text[i+2]
            l = re.findall('\d*\.?\d+', sent)
            for i in l:
                a = float(i)
            return (round(a,2))
            #return (sent)
        
        
# Function to extract skill set details
def skillSet(Text):
    t = []
    for i in range(len(Text)):
        if Text[i] in allTerms:
            if Text[i] in t:
                continue
            t.append(Text[i]) 
    return(t)

html_temp = """
    <div style="background-color:orange;padding:7px">
    <h2 style="color:white;text-align:center;"> Resume Classification App</h2>
    </div>
    """
html_temp1 = """
        <div style="background-color:green;padding:5px">
        <h2 style="color:white;text-align:center;">RESULTS </h2>
        </div>
        """
# this is the main function in which we define our webpage  
def main():

    st.markdown(html_temp,unsafe_allow_html=True)
    path = st.text_input('Enter the resumes folder path')
    st.markdown(html_temp1,unsafe_allow_html=True)
    st.subheader(' ')

    # following lines create boxes in which user can enter data required to make prediction 
    #path ='D:/Python_Practice/P-130_Project/Resumes'
    #path = st.sidebar.text_input('Enter the resumes folder path')

    # Create an empty Data Frame ResumeText with two columns
    ResumeText = pd.DataFrame([], columns=['Name', 'Exp_years', 'SkillSet','RoleApplied'])
    
    # when 'Predict' is clicked, make the prediction and store it 
    #if st.button("Upload and Get Result"): 
        # Search the directory path and loop through the resume documents and callthe function getText
    #if st.button("Process"):
    for filename in os.listdir(path):
        filename = os.path.join(path, filename)
        extText = getText(filename)
        tokText = tokenText(extText)
        role = roleApplied(tokText)
        Name = contactDetails(tokText)
        experience = expDetails(extText)
        skills = skillSet(tokText)
        NewRow = [Name,experience, skills,role]
        ResumeText.loc[len(ResumeText)] = NewRow
        #st.dataframe(ResumeText)
        java = (ResumeText["RoleApplied"] == "Java Developer")
        #javares = ResumeText[java]
        workday = (ResumeText["RoleApplied"] == "WorkDay ERP")
        peosoft = (ResumeText["RoleApplied"] == "Peoplesoft")
        dbms = (ResumeText["RoleApplied"] == "Database Developer")
    
    html_temp10 = """
        <div style="background-color:Blue;padding:5px">
        <h2 style="color:white;text-align:center;">tab1 </h2>
        </div>
        """
    tab1, tab2, tab3, tab4, tab5  = st.tabs(["Overview", "JAVA", "DBMS","Peoplesoft","Workday"])



    with tab1:
        st.subheader("No of Resumes Received")
        num = pd.DataFrame(ResumeText['RoleApplied'].value_counts())
        num['Category'] = num.index
        num.set_axis(['No of Resumes', 'Category'], axis='columns', inplace=True)
        num.reset_index(inplace=True, drop=True)
        num = num[['Category', 'No of Resumes']]
        st.dataframe(num)

        c = alt.Chart(num).mark_bar().encode(
            x='Category',
            y='No of Resumes')
        st.altair_chart(c, use_container_width=True)
        

        #st.dataframe(ResumeText)
    with tab2:
        st.subheader("JAVA")
        st.dataframe(ResumeText[java])
    with tab3:
        st.subheader("DBMS")
        st.dataframe(ResumeText[dbms])
    with tab4:
        st.subheader("Peoplesoft")
        st.dataframe(ResumeText[peosoft])
    with tab5:
        st.subheader("Workday")
        st.dataframe(ResumeText[workday])


if __name__=='__main__':
    main()

